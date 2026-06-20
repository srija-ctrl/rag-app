import io
import re
import time
import httpx
from typing import List, Tuple

# PDF
import pypdf

# DOCX
from docx import Document as DocxDocument

CHUNK_SIZE = 800     # characters per chunk
CHUNK_OVERLAP = 150  # overlap between chunks


# ── Text splitter ──────────────────────────────────────────────────────────────

def split_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Split text into overlapping chunks, respecting sentence/paragraph boundaries."""
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(text) <= chunk_size:
        return [text] if text else []

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        if end >= len(text):
            chunks.append(text[start:].strip())
            break
        # Try to break at paragraph, then sentence, then word
        for sep in ["\n\n", ". ", "? ", "! ", " "]:
            pos = text.rfind(sep, start + chunk_size // 2, end)
            if pos != -1:
                end = pos + len(sep)
                break
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - overlap
    return [c for c in chunks if len(c) > 30]


# ── Extractors ─────────────────────────────────────────────────────────────────

def extract_pdf(data: bytes) -> str:
    reader = pypdf.PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        t = page.extract_text() or ""
        if t.strip():
            pages.append(t)
    return "\n\n".join(pages)


def extract_docx(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also grab table cells
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n\n".join(paragraphs)


def extract_text(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


async def extract_url(url: str) -> Tuple[str, str]:
    """Fetch URL and extract readable text. Returns (title, text)."""
    async with httpx.AsyncClient(
        timeout=30,
        follow_redirects=True,
        headers={"User-Agent": "Mozilla/5.0 (compatible; RAGBot/1.0)"},
    ) as client:
        resp = await client.get(url)
        resp.raise_for_status()
        html = resp.text

    # Simple HTML text extraction without extra deps
    # Strip scripts, styles, tags
    html = re.sub(r"<script[^>]*>.*?</script>", " ", html, flags=re.DOTALL | re.IGNORECASE)
    html = re.sub(r"<style[^>]*>.*?</style>", " ", html, flags=re.DOTALL | re.IGNORECASE)

    title_match = re.search(r"<title[^>]*>(.*?)</title>", html, re.IGNORECASE | re.DOTALL)
    title = re.sub(r"<[^>]+>", "", title_match.group(1)).strip() if title_match else url

    # Convert block tags to newlines
    html = re.sub(r"<(p|div|h[1-6]|li|br|tr)[^>]*>", "\n", html, flags=re.IGNORECASE)
    # Strip remaining tags
    text = re.sub(r"<[^>]+>", " ", html)
    # Decode common entities
    for ent, ch in [("&amp;", "&"), ("&lt;", "<"), ("&gt;", ">"), ("&nbsp;", " "), ("&#39;", "'"), ("&quot;", '"')]:
        text = text.replace(ent, ch)
    text = re.sub(r" {2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return title, text.strip()
